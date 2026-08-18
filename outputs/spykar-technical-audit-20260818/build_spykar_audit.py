from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'outputs/spykar-technical-audit-20260818/Spykar_Technical_SEO_Audit_2026-08-18.docx'

def shade(cell, fill):
    p = cell._tc.get_or_add_tcPr(); s = OxmlElement('w:shd'); s.set(qn('w:fill'), fill); p.append(s)

def text(cell, value, bold=False, size=8):
    cell.text = ''
    p = cell.paragraphs[0]; r = p.add_run(str(value)); r.bold = bold; r.font.size = Pt(size)

def heading(doc, value, level=1):
    p = doc.add_heading(value, level=level)
    for r in p.runs: r.font.color.rgb = RGBColor(23,58,87)
    return p

def bullets(doc, rows):
    for row in rows:
        doc.add_paragraph(row, style='List Bullet')

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,'173A57'); text(c,h,True,8)
        for r in c.paragraphs[0].runs: r.font.color.rgb=RGBColor(255,255,255)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            shade(cells[i],'EEF3F7' if len(t.rows)%2==0 else 'FFFFFF'); text(cells[i],v,False,7)
            cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    return t

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.55); sec.right_margin=Inches(.55)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Spykar Technical SEO Audit'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor(23,58,87)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Evidence period: 17 Feb–15 Aug 2026 vs 21 Aug 2025–16 Feb 2026 | Run date: 18 Aug 2026').italic=True
heading(doc,'Executive summary',1)
doc.add_paragraph('Organic visibility and Organic Search traffic have declined materially even as average Google position improved. This points to a mix of lower search demand/impressions, weak SERP click-through rate, and page-level commercial/content gaps—not a single confirmed crawling failure. The highest-impact work is to improve category-page SERP appeal and intent coverage, resolve GA4 landing-page attribution, and verify sitemap/indexing coverage with URL Inspection.')
table(doc,['Metric','Current','Previous','Change'],[
['GSC clicks','140,387','283,684','-50.5%'],['GSC impressions','11.66M','19.19M','-39.2%'],['GSC CTR','1.20%','1.48%','-18.5% relative'],['GSC average position','9.35','11.65','Improved 2.30 positions'],['GA4 Organic sessions','161,731','317,931','-49.1%'],['GA4 Organic revenue','₹5.19M','₹19.22M','-73.0%'],['GA4 Organic purchases','2,333','9,025','-74.2%']])
heading(doc,'Priority findings and action register',1)
table(doc,['Category','Issue','Data source','Affected URLs','Current data','SEO impact','Priority','Recommended fix'],[
['Organic performance','Visibility and traffic decline despite better rankings','GSC + GA4','Sitewide; especially / and /collections/men-jeans','Clicks -50.5%; sessions -49.1%; position improved 11.65→9.35','Major demand/CTR and revenue risk','Critical','Segment queries/pages by brand vs non-brand and device; refresh priority commercial hubs first; measure weekly impressions, CTR, sessions and revenue.'],
['CTR opportunity','High-impression terms rank 4–10 with extremely low CTR','GSC','Relevant category hubs','“trousers for men”: 137,285 impr., pos. 6.78, 0.40% CTR; “black jeans for women”: 120,915, pos. 7.66, 0.10% CTR','Large low-effort click upside','High','Map each term to one canonical category, rewrite title/meta/H1 and strengthen category copy, product assortment cues and internal links; test SERP snippets.'],
['Ranking opportunity','High-volume commercial terms remain outside top 10','GSC','Women jeans, cargo, men jeans hubs','“jeans for women” pos. 13.48; “jeans for men” pos. 12.95; “cargo pants for men” pos. 10.64','Recoverable non-brand demand','High','Expand corresponding collection pages around fit, style, material and use-case intent; add contextual links from relevant blogs/navigation.'],
['Commercial landing pages','Men Jeans traffic loss is commercially material','GSC + GA4','/collections/men-jeans','GSC clicks -21,929; GA4 sessions -20,533; 203 purchases / ₹503k current period','High-revenue category loss','High','Run page/query mapping, compare lost terms, check filters/canonical/stock availability, then refresh on-page copy and internal-link prominence.'],
['Measurement','GA4 landing page attribution gap','GA4','(not set)','3,895 Organic sessions; 1.6% engagement; zero purchases','Obscures SEO landing-page diagnosis','High','Audit GA4 page_view/landing-page capture, consent mode and server-side/client-side tagging; validate in DebugView before reporting page-level conversion rates.'],
['Sitemap/indexation','Sitemap reports zero indexed URLs in the returned sitemap row','GSC + live sitemap','https://spykar.com/sitemap.xml','4,872 web URLs submitted; 0 indexed; 0 errors/warnings; index downloaded 15 Aug','Potential reporting or index-coverage issue','High','Do not assume deindexation: inspect representative home, category, product and blog URLs in GSC; compare Page Indexing totals and child-sitemap data.'],
['Metadata / duplicate risk','Repeated generic category descriptions visible in URL inventory','Firecrawl URL inventory','Multiple collection pages','Several collection pages share the same generic “Welcome to our world…” description','Weak relevance and CTR differentiation','Medium','Crawl all indexable collection pages; deduplicate titles/descriptions and prioritise pages with GSC impressions.'],
['Site architecture','Duplicate/near-duplicate route candidates need canonical review','Firecrawl URL inventory','/pages/faq and /pages/faqs; /a/storelocator and /pages/store-locator','Both route variants discovered','Possible duplication/canonical ambiguity','Medium','Check final status, canonical, internal-link targets and GSC index state; retain one canonical route per intent and 301 obsolete duplicates when appropriate.'],
['Conversion','High-traffic commercial pages record no purchases','GA4','/collections/men-trousers-joggers; /collections/check-shirt-for-men','1,595 and 1,222 Organic sessions respectively; 0 purchases','Commercial leakage or tracking gap','Medium','QA availability, filtering, price/promotion, add-to-cart and event capture; compare against product-list conversion baseline.']])
heading(doc,'Validated technical evidence',1)
bullets(doc,[
'Robots.txt returns 200, permits public products/collections/pages/blogs, blocks transactional paths, and declares https://spykar.com/sitemap.xml.',
'The live sitemap index returns 200 and references product, page, collection, blog and agentic-discovery sitemaps.',
'Native Firecrawl rendered fetches returned HTTP 200 for the tested home, category, blog and About pages. The homepage metadata is present and the rendered HTML exposes navigational internal links.',
'Browser-tool and Composio Firecrawl routes were unavailable; therefore, no claim is made about sitewide broken links, redirect chains, complete crawl depth, mobile visual defects, console errors, or exhaustive schema validation.'])
heading(doc,'Quick wins: 0–30 days',1)
bullets(doc,['Refresh Men Jeans, Men Shirts, Men T-shirts, Women Jeans, Men Trousers/Joggers and Cargo category titles, descriptions, H1s and above-the-fold copy against their mapped query clusters.','Fix GA4 landing-page attribution before using page-level conversion findings for investment decisions.','Use GSC URL Inspection for the homepage, men-jeans, women-jeans, a representative product and a representative blog; document canonical, crawl and index verdicts.','Create a duplicate-route register for FAQ and store-locator variants, then correct internal links and canonicals/redirects after validation.'])
heading(doc,'30–60 and 60–90 day plan',1)
bullets(doc,['30–60 days: establish category-to-query ownership; prevent cannibalisation; add curated editorial links into priority collections; assess product availability and filter/index rules.','60–90 days: complete a full Screaming Frog or equivalent crawl, validate Core Web Vitals by mobile/desktop field data, audit structured data and expand category-supporting content where demand warrants.'])
heading(doc,'Scope and evidence notes',1)
doc.add_paragraph('This is a partial technical audit using native Firecrawl, GSC and GA4. GSC uses the verified https://spykar.com/ property; GA4 uses www.spykar.com - GA4 (properties/320857986). The audit does not claim a sitewide crawl pass because the configured brokered crawl/browser services were unavailable. All recommendations require post-release verification.')
doc.save(OUT)
