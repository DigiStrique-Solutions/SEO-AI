from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[4]
SOURCE = ROOT / "brands/offlimits/blogs/swat-collection-performance-street-style.md"
OUTPUT = ROOT / "brands/offlimits/blogs/drafts/offlimits-swat-collection-unsanitized.docx"


def set_font(run, size=11, bold=None, italic=None, color="000000", underline=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "DADCE0")
        borders.append(tag)


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url, bold=False, italic=False):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    rPr.append(size)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


TOKEN_RE = re.compile(r"(\[\*\*.+?\*\*\]\(.+?\)|\[.+?\]\(.+?\)|\*\*.+?\*\*|\*.+?\*)")


def add_inline(paragraph, text):
    text = text.replace("—", "-").replace("–", "-")
    for token in TOKEN_RE.split(text):
        if not token:
            continue
        m = re.fullmatch(r"\[\*\*(.+?)\*\*\]\((.+?)\)", token)
        if m:
            add_hyperlink(paragraph, m.group(1), m.group(2), bold=True)
            continue
        m = re.fullmatch(r"\[(.+?)\]\((.+?)\)", token)
        if m:
            add_hyperlink(paragraph, m.group(1), m.group(2))
            continue
        m = re.fullmatch(r"\*\*(.+?)\*\*", token)
        if m:
            run = paragraph.add_run(m.group(1))
            set_font(run, bold=True)
            continue
        m = re.fullmatch(r"\*(.+?)\*", token)
        if m:
            run = paragraph.add_run(m.group(1))
            set_font(run, italic=True, color="555555")
            continue
        run = paragraph.add_run(token)
        set_font(run)


def style_paragraph(paragraph, after=8, line=1.15):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_image(doc, rel_path, alt_text):
    path = SOURCE.parent / rel_path
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=Inches(6.5))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)


def make_doc():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after, color in (
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = lines[0].removeprefix("# ")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, size=26, bold=False)

    status = doc.add_paragraph()
    status.paragraph_format.space_after = Pt(12)
    sr = status.add_run("Editorial status: Draft - live ZeroGPT verification pending (provider returned 'Not enough credits').")
    set_font(sr, size=9.5, italic=True, color="555555")

    i = 1
    while i < len(lines):
        line = lines[i]
        if not line.strip() or line.strip() == "---":
            i += 1
            continue
        if line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            add_image(doc, m.group(2), m.group(1))
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:])
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:])
            i += 1
            continue
        if line.startswith("| "):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"-+", c) for c in cells):
                    rows.append(cells)
                i += 1
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            widths = [Inches(1.35), Inches(2.05), Inches(1.75), Inches(1.35)]
            for ridx, row in enumerate(rows):
                for cidx, value in enumerate(row):
                    cell = table.cell(ridx, cidx)
                    cell.width = widths[cidx]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell)
                    cp = cell.paragraphs[0]
                    style_paragraph(cp, after=0, line=1.0)
                    add_inline(cp, value)
                    for run in cp.runs:
                        set_font(run, size=9.5, bold=(ridx == 0))
            set_table_borders(table)
            set_table_geometry(table, [1944, 2952, 2520, 1944])
            table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            style_paragraph(p, after=4)
            add_inline(p, line[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        style_paragraph(p)
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line.strip("*"))
            set_font(run, size=9, italic=True, color="555555")
        else:
            add_inline(p, line)
        i += 1

    props = doc.core_properties
    props.title = title
    props.subject = "OFFLIMITS S.W.A.T. collection blog draft"
    props.author = "Strique SEO"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    make_doc()
